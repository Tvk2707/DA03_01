from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = "Giai_thich_toan_bo_code_DA03_01.docx"


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=16, bold=True, color="2E74B5")
    elif level == 2:
        set_font(r, size=13, bold=True, color="2E74B5")
    else:
        set_font(r, size=12, bold=True, color="1F4D78")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    r = p.add_run(text)
    set_font(r)
    return p


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, name="Consolas", size=9, color="333333")
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], "E8EEF5")
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, bold=True, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=10)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def footer(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("DA03_01 - Giai thich code cho nguoi moi hoc")
        set_font(r, size=9, color="666666")


def main():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Giải thích toàn bộ code dự án DA03_01")
    set_font(r, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Bản giải thích dễ hiểu cho dự án Java Web JSP/Servlet + SQL Server")
    set_font(r, size=12, color="555555")

    add_body(
        doc,
        "Tài liệu này giải thích ý nghĩa các phần code chính trong dự án. Mục tiêu không phải học thuộc từng dòng, "
        "mà là hiểu dự án chạy theo luồng nào, mỗi thư mục chịu trách nhiệm gì, và khi sửa lỗi hoặc thêm chức năng thì nên đọc file nào trước.",
    )

    add_heading(doc, "1. Bức tranh tổng quát", 1)
    add_body(
        doc,
        "Đây là một ứng dụng Java Web đóng gói dạng WAR. Giao diện dùng JSP, CSS và JavaScript. Backend dùng Servlet, Service, DAO. "
        "Dữ liệu lấy từ SQL Server thông qua JDBC; các Entity JPA/Hibernate cũng được khai báo để mô tả bảng trong database.",
    )
    add_bullet(doc, "FE là phần người dùng nhìn thấy: JSP, CSS, JavaScript, ảnh sản phẩm.")
    add_bullet(doc, "BE là phần xử lý phía server: Controller nhận request, Service quyết định nghiệp vụ, DAO truy vấn database.")
    add_bullet(doc, "Database là SQL Server tên quan_ly_ban_kinh.")

    add_heading(doc, "Luồng chạy dễ nhớ", 2)
    for line in [
        "Người dùng mở trang hoặc bấm nút trên JSP",
        "Servlet Controller nhận request theo URL",
        "Controller gọi Service",
        "Service gọi DAO",
        "DAO chạy SQL qua JDBC và lấy dữ liệu từ SQL Server",
        "Dữ liệu được đưa vào Model/View object",
        "Controller forward sang JSP để hiển thị lại cho người dùng",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "2. Cấu trúc thư mục", 1)
    add_table(
        doc,
        ["Vị trí", "Ý nghĩa"],
        [
            ["pom.xml", "Khai báo dự án Maven, đóng gói WAR, thư viện Servlet, Hibernate, SQL Server JDBC, Lombok, JUnit."],
            ["src/main/java/BE/controller", "Các Servlet nhận request từ trình duyệt. Ví dụ: /admin/hoa-don và /admin/hoa-don/chi-tiet."],
            ["src/main/java/BE/service", "Tầng trung gian xử lý nghiệp vụ. Controller không gọi SQL trực tiếp mà đi qua Service."],
            ["src/main/java/BE/dao", "Tầng làm việc trực tiếp với SQL Server. Chứa câu SQL SELECT/INSERT/UPDATE."],
            ["src/main/java/BE/Model", "Các object dùng để đưa dữ liệu lên màn hình hoặc nhận dữ liệu từ form."],
            ["src/main/java/BE/Entity", "Các class ánh xạ bảng database bằng JPA/Hibernate."],
            ["src/main/java/BE/jdbc", "Quản lý kết nối SQL Server và file test kết nối database."],
            ["src/main/resources/META-INF/persistence.xml", "Cấu hình JPA/Hibernate: driver, URL database, user, password, dialect."],
            ["src/main/webapp/FE/Admin", "Frontend admin: JSP, layout, CSS, JavaScript, ảnh sản phẩm."],
            ["src/main/webapp/WEB-INF/web.xml", "Cấu hình web app, hiện đang đặt index.jsp là trang welcome."],
        ],
        [2600, 6760],
    )

    add_heading(doc, "3. Các file cấu hình quan trọng", 1)
    add_heading(doc, "pom.xml", 2)
    add_body(doc, "File này nói cho Maven biết dự án là Java Web và cần những thư viện nào.")
    add_bullet(doc, "packaging = war: khi build sẽ ra file WAR để chạy trên server như Tomcat.")
    add_bullet(doc, "jakarta.servlet-api: dùng để viết Servlet/JSP.")
    add_bullet(doc, "hibernate-core: dùng cho JPA/Hibernate Entity.")
    add_bullet(doc, "mssql-jdbc: driver để Java kết nối SQL Server.")
    add_bullet(doc, "maven-war-plugin: hỗ trợ đóng gói ứng dụng web.")

    add_heading(doc, "web.xml và index.jsp", 2)
    add_body(doc, "web.xml đặt index.jsp làm trang đầu tiên. index.jsp hiện chuyển hướng người dùng sang trang thống kê admin.")
    add_code_line(doc, "index.jsp -> /FE/Admin/Thongke.jsp")

    add_heading(doc, "persistence.xml", 2)
    add_body(
        doc,
        "File này cấu hình JPA/Hibernate kết nối tới SQL Server. Database đang trỏ tới localhost:1433, database quan_ly_ban_kinh, user minh.",
    )
    add_bullet(doc, "hibernate.hbm2ddl.auto = validate: Hibernate chỉ kiểm tra mapping có khớp bảng hay không, không tự tạo bảng.")
    add_bullet(doc, "hibernate.show_sql = true: khi chạy JPA có thể in câu SQL ra console.")

    add_heading(doc, "4. Backend hoạt động như thế nào", 1)
    add_heading(doc, "Controller", 2)
    add_body(doc, "Controller là cửa vào của backend. Nó nhận request từ trình duyệt và quyết định gọi Service nào.")
    add_table(
        doc,
        ["File", "URL", "Nhiệm vụ"],
        [
            ["HoaDonController.java", "/admin/hoa-don", "Hiển thị danh sách hoá đơn, nhận form thêm/sửa hoá đơn, đổi trạng thái."],
            ["ChiTietHoaDonController.java", "/admin/hoa-don/chi-tiet", "Hiển thị chi tiết một hoá đơn: sản phẩm, thanh toán, lịch sử."],
        ],
        [2500, 2200, 4660],
    )

    add_heading(doc, "Service", 2)
    add_body(
        doc,
        "HoaDonService.java là tầng trung gian. Nó giúp Controller không phải biết SQL. Service gọi DAO, đồng thời xử lý một số luật đơn giản như tạo mã hoá đơn nếu mã bị trống hoặc bị trùng.",
    )
    add_bullet(doc, "getAllHoaDon(): lấy danh sách hoá đơn.")
    add_bullet(doc, "getHoaDonById(id): lấy một hoá đơn theo id.")
    add_bullet(doc, "saveHoaDon(...): nếu chưa có id thì thêm mới, nếu có id thì cập nhật.")
    add_bullet(doc, "updateTrangThai(...): đổi trạng thái và ghi chú.")

    add_heading(doc, "DAO", 2)
    add_body(
        doc,
        "HoaDonDAO.java là nơi viết SQL thật. Nếu dữ liệu trên màn hình không đúng, đây thường là file cần kiểm tra sớm vì nó quyết định lấy bảng nào, join bảng nào, và map cột SQL sang object Java ra sao.",
    )
    add_table(
        doc,
        ["Nhóm", "Method", "Ý nghĩa"],
        [
            ["READ", "findAll()", "Lấy danh sách hoá đơn, join nhân viên, khách hàng, phiếu giảm giá."],
            ["READ", "findById(id)", "Lấy thông tin một hoá đơn theo id."],
            ["READ", "findAllNhanVien()", "Lấy danh sách nhân viên để đưa vào combobox khi thêm hoá đơn."],
            ["READ", "findAllSanPhamHoaDon()", "Lấy sản phẩm còn hàng để chọn khi tạo hoá đơn."],
            ["READ", "findDetailsByHoaDonId(id)", "Lấy danh sách sản phẩm trong một hoá đơn."],
            ["READ", "findPaymentsByHoaDonId(id)", "Lấy các lần thanh toán của hoá đơn."],
            ["READ", "findPaymentHistoryByHoaDonId(id)", "Lấy lịch sử thanh toán hệ thống."],
            ["READ", "findHistoryByHoaDonId(id)", "Lấy lịch sử xử lý hoá đơn."],
            ["CREATE", "insert(hoaDon)", "Thêm một hoá đơn mới vào bảng hoa_don."],
            ["CREATE", "insertChiTietHoaDon(...)", "Thêm sản phẩm vào chi_tiet_hoa_don và tính lại tổng tiền."],
            ["UPDATE", "update(hoaDon)", "Sửa thông tin cơ bản của hoá đơn."],
            ["UPDATE", "updateStatus(...)", "Đổi trạng thái hoá đơn và ghi vào lich_su_hoa_don."],
            ["DELETE mềm", "delete(id)", "Không xoá dòng khỏi database, chỉ đổi trạng thái thành 5 để ẩn/huỷ."],
        ],
        [1400, 2700, 5260],
    )

    add_heading(doc, "JDBC và JPA", 2)
    add_body(doc, "Dự án có cả JDBC và JPA. Phần hoá đơn hiện đang dùng JDBC trực tiếp qua HoaDonDAO.")
    add_bullet(doc, "DatabaseConnectionManager.java tạo chuỗi kết nối SQL Server, đọc cấu hình từ biến môi trường nếu có.")
    add_bullet(doc, "JdbcMain.java là file chạy thử kết nối database và EntityManager.")
    add_bullet(doc, "EntityManagerUtils.java tạo EntityManagerFactory cho JPA/Hibernate.")

    add_heading(doc, "5. Model và Entity khác nhau thế nào", 1)
    add_body(
        doc,
        "Entity là class đại diện bảng trong database. Model/View object là class gọn hơn, thường dùng để đưa đúng dữ liệu cần hiển thị lên JSP. Vì vậy không phải màn hình nào cũng cần đưa nguyên Entity lên.",
    )
    add_table(
        doc,
        ["Loại", "Ví dụ", "Dùng để làm gì"],
        [
            ["Entity", "HoaDon, SanPham, KhachHang, NhanVien", "Ánh xạ bảng thật trong SQL Server bằng @Entity và @Table."],
            ["Model/View", "HoaDonView, ChiTietHoaDonView", "Chứa dữ liệu đã join/đã format theo nhu cầu màn hình."],
            ["Input", "ChiTietHoaDonInput", "Nhận dữ liệu người dùng chọn từ form: id sản phẩm chi tiết và số lượng."],
        ],
        [1800, 3000, 4560],
    )

    add_heading(doc, "6. Frontend hoạt động như thế nào", 1)
    add_body(
        doc,
        "Frontend nằm dưới src/main/webapp/FE/Admin. JSP vừa chứa HTML vừa đọc dữ liệu request attribute do Controller truyền sang. CSS quyết định giao diện. JavaScript xử lý tương tác trên trình duyệt.",
    )
    add_table(
        doc,
        ["File", "Ý nghĩa"],
        [
            ["FE/Admin/layout/sidebar.jsp", "Menu bên trái dùng chung cho các trang admin."],
            ["FE/Admin/layout/header.jsp", "Header phía trên dùng chung."],
            ["FE/Admin/Thongke.jsp", "Trang thống kê được index.jsp chuyển hướng tới."],
            ["FE/Admin/QuanLyHoaDon/quan_ly_hoa_don.jsp", "Trang danh sách hoá đơn: bộ lọc, thống kê nhanh, form thêm/sửa, bảng danh sách."],
            ["FE/Admin/QuanLyHoaDon/chi_tiet_hoa_don.jsp", "Trang chi tiết hoá đơn: sản phẩm, thanh toán, trạng thái, lịch sử."],
            ["FE/Admin/QuanLyHoaDon/hoa_don.js", "Lọc bảng, mở/đóng form, tính tổng tiền theo sản phẩm, xuất CSV, in hoá đơn, mở modal."],
            ["FE/Admin/css/hoa_don.css", "Toàn bộ style riêng cho màn hình hoá đơn."],
            ["FE/Admin/QuanLyMaGiamGia/*.jsp/js", "Nhóm giao diện quản lý mã/đợt giảm giá."],
        ],
        [3600, 5760],
    )

    add_heading(doc, "7. Luồng hoá đơn cụ thể", 1)
    add_heading(doc, "Mở danh sách hoá đơn", 2)
    for line in [
        "Trình duyệt gọi /admin/hoa-don.",
        "HoaDonController.doGet() chạy.",
        "Controller gọi HoaDonService.getAllHoaDon(), getAllNhanVien(), getAllSanPhamHoaDon().",
        "Service gọi HoaDonDAO.",
        "DAO SELECT dữ liệu từ SQL Server.",
        "Controller setAttribute rồi forward sang quan_ly_hoa_don.jsp.",
        "JSP dùng hoaDonList để vẽ bảng hoá đơn.",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "Thêm hoá đơn", 2)
    for line in [
        "Người dùng bấm Thêm hoá đơn trên JSP.",
        "hoa_don.js mở form và tính tổng tiền khi chọn sản phẩm/số lượng.",
        "Form gửi POST về /admin/hoa-don với action=save.",
        "HoaDonController.doPost() đọc dữ liệu form thành HoaDonView và danh sách ChiTietHoaDonInput.",
        "HoaDonService.saveHoaDon() quyết định thêm mới hay cập nhật.",
        "HoaDonDAO.insert() thêm vào bảng hoa_don.",
        "HoaDonDAO.insertChiTietHoaDon() thêm sản phẩm vào chi_tiet_hoa_don và cập nhật tổng tiền.",
        "Controller redirect về /admin/hoa-don để tải lại danh sách.",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "Xem chi tiết hoá đơn", 2)
    for line in [
        "Người dùng bấm icon mắt ở bảng danh sách.",
        "Trình duyệt mở /admin/hoa-don/chi-tiet?id=...",
        "ChiTietHoaDonController lấy id từ URL.",
        "Controller lấy hoá đơn, sản phẩm, thanh toán, lịch sử xử lý.",
        "Dữ liệu được forward sang chi_tiet_hoa_don.jsp.",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "Đổi trạng thái hoặc huỷ hoá đơn", 2)
    add_body(
        doc,
        "Các nút thanh toán, đổi trạng thái, huỷ đơn đều gửi POST về /admin/hoa-don với action=changeStatus. DAO cập nhật bảng hoa_don và ghi thêm dòng vào lich_su_hoa_don để biết hoá đơn đã được xử lý lúc nào.",
    )

    add_heading(doc, "8. Các bảng database liên quan phần hoá đơn", 1)
    add_table(
        doc,
        ["Bảng", "Vai trò trong code"],
        [
            ["hoa_don", "Bảng chính lưu mã hoá đơn, người nhận, tổng tiền, trạng thái, ghi chú."],
            ["chi_tiet_hoa_don", "Lưu từng sản phẩm nằm trong hoá đơn."],
            ["san_pham_chi_tiet", "Nguồn sản phẩm, giá bán, tồn kho để chọn khi tạo hoá đơn."],
            ["san_pham, danh_muc, thuong_hieu, chat_lieu, ...", "Các bảng mô tả thông tin sản phẩm để hiển thị đẹp ở chi tiết hoá đơn."],
            ["nhan_vien", "Nhân viên xử lý hoá đơn."],
            ["khach_hang", "Khách hàng gắn với hoá đơn nếu có."],
            ["phieu_giam_gia", "Voucher/mã giảm giá join vào danh sách hoá đơn."],
            ["thanh_toan_hoa_don", "Các lần thanh toán thực tế của hoá đơn."],
            ["lich_su_thanh_toan", "Lịch sử thanh toán hệ thống."],
            ["lich_su_hoa_don", "Lịch sử hành động với hoá đơn: cập nhật trạng thái, huỷ, xoá mềm."],
        ],
        [3000, 6360],
    )

    add_heading(doc, "9. Cách đọc code khi mới học", 1)
    add_body(doc, "Nếu bạn mới học, đừng mở tất cả file cùng lúc. Hãy đọc theo thứ tự sau:")
    for line in [
        "Bước 1: Mở JSP để biết màn hình có những nút/form/table nào.",
        "Bước 2: Nhìn action hoặc href trong JSP để biết URL gửi về đâu.",
        "Bước 3: Tìm @WebServlet tương ứng trong controller.",
        "Bước 4: Từ controller đi xuống service.",
        "Bước 5: Từ service đi xuống DAO để xem SQL.",
        "Bước 6: Nhìn Model/View để biết dữ liệu SQL được đưa lên màn hình bằng field nào.",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "10. Ghi chú quan trọng", 1)
    add_bullet(doc, "Một số chữ tiếng Việt trong source đang bị lỗi mã hoá hiển thị. Logic vẫn có thể chạy, nhưng comment và text UI nên được sửa lại encoding sau.")
    add_bullet(doc, "Một vài Entity đang tham chiếu package com.eyewear.entity trong field quan hệ, trong khi file hiện nằm ở package BE.Entity. Nếu dùng sâu JPA, phần này cần kiểm tra lại.")
    add_bullet(doc, "Phần hoá đơn dùng JDBC trực tiếp nên các lỗi dữ liệu thường nằm ở SQL trong HoaDonDAO hoặc cấu hình kết nối trong DatabaseConnectionManager.")
    add_bullet(doc, "Xoá hoá đơn hiện là xoá mềm: đổi trạng thái thay vì xoá dòng khỏi database. Cách này an toàn hơn vì hoá đơn còn liên kết với chi tiết, thanh toán, lịch sử.")

    add_heading(doc, "11. Tóm tắt ngắn gọn", 1)
    add_body(
        doc,
        "Dự án này có thể hiểu đơn giản là: JSP hiển thị giao diện, Servlet nhận request, Service xử lý nghiệp vụ, DAO truy vấn SQL Server, Model/View chuyển dữ liệu sang JSP, Entity mô tả bảng database. "
        "Khi cần sửa chức năng hoá đơn, hãy đi theo đường: quan_ly_hoa_don.jsp hoặc chi_tiet_hoa_don.jsp -> HoaDonController/ChiTietHoaDonController -> HoaDonService -> HoaDonDAO -> SQL Server.",
    )

    footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    main()
