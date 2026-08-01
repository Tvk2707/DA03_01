from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path(r"D:\Java3\DA1_03\Bảng 3.1 - Danh sách bảng RIOR.docx")


ROWS = [
    (1, "thuong_hieu", "Lưu danh mục thương hiệu của sản phẩm kính.", ""),
    (2, "danh_muc", "Lưu danh mục phân loại sản phẩm kính.", ""),
    (3, "chat_lieu", "Lưu danh mục chất liệu của sản phẩm.", ""),
    (4, "kieu_dang", "Lưu danh mục kiểu dáng của sản phẩm kính.", ""),
    (5, "mau_sac", "Lưu danh mục màu sắc của sản phẩm.", ""),
    (6, "kich_co", "Lưu danh mục kích cỡ của sản phẩm.", ""),
    (7, "trong_kinh", "Lưu thông tin các loại tròng kính.", ""),
    (8, "hinh_dang_gong", "Lưu danh mục hình dáng gọng kính.", ""),
    (9, "kieu_quai_kinh", "Lưu danh mục kiểu quai kính.", ""),
    (10, "gong_kinh", "Lưu thông tin gọng kính theo hình dáng và kiểu quai.", "hinh_dang_gong, kieu_quai_kinh"),
    (11, "san_pham", "Lưu thông tin chung và các thuộc tính phân loại của sản phẩm kính.", "danh_muc, thuong_hieu, chat_lieu, kieu_dang, gong_kinh, trong_kinh"),
    (12, "san_pham_chi_tiet", "Lưu các biến thể sản phẩm theo màu sắc, kích cỡ, giá bán và số lượng tồn kho.", "san_pham, mau_sac, kich_co"),
    (13, "hinh_anh_san_pham", "Lưu hình ảnh của từng sản phẩm.", "san_pham"),
    (14, "khach_hang", "Lưu hồ sơ và thông tin liên hệ của khách hàng.", ""),
    (15, "dia_chi_khach_hang", "Lưu các địa chỉ gắn với khách hàng.", "khach_hang"),
    (16, "nhan_vien", "Lưu hồ sơ, tài khoản đăng nhập, vai trò và trạng thái của nhân viên.", ""),
    (17, "ca_lam_viec", "Lưu thông tin ca làm việc của nhân viên.", "nhan_vien"),
    (18, "phieu_giam_gia", "Lưu thông tin, điều kiện, thời hạn và số lượng phiếu giảm giá.", ""),
    (19, "khach_hang_phieu_giam_gia", "Lưu thông tin phiếu giảm giá được cấp cho từng khách hàng.", "khach_hang, phieu_giam_gia"),
    (20, "hoa_don", "Lưu thông tin giao dịch bán hàng, tổng tiền và trạng thái hóa đơn.", "khach_hang, nhan_vien, phieu_giam_gia, ca_lam_viec"),
    (21, "chi_tiet_hoa_don", "Lưu sản phẩm, số lượng và đơn giá trong từng hóa đơn.", "hoa_don, san_pham_chi_tiet"),
    (22, "hinh_thuc_thanh_toan", "Lưu danh mục các hình thức thanh toán.", ""),
    (23, "thanh_toan_hoa_don", "Lưu số tiền, hình thức, mã giao dịch và thông tin thanh toán hóa đơn.", "hoa_don, hinh_thuc_thanh_toan"),
    (24, "lich_su_hoa_don", "Lưu lịch sử thay đổi trạng thái và xử lý hóa đơn.", "hoa_don"),
    (25, "lich_su_thanh_toan", "Lưu lịch sử các sự kiện thanh toán của hóa đơn.", "hoa_don"),
]


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def add_cell_text(cell, text, bold=False, center=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    set_run_font(run, size=12, bold=bold)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(5)
title.paragraph_format.keep_with_next = True
title_run = title.add_run("Bảng 3.1: Danh sách bảng")
set_run_font(title_run, size=13, bold=True, italic=True)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.allow_autofit = False
set_borders(table)

headers = ["STT", "Tên bảng", "Mô tả", "Phụ thuộc"]
for idx, label in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=100, start=110, bottom=100, end=110)
    add_cell_text(cell, label, bold=True, center=(idx == 0))
set_repeat_table_header(table.rows[0])
prevent_row_split(table.rows[0])

for stt, name, description, dependency in ROWS:
    cells = table.add_row().cells
    values = [stt, name, description, dependency]
    for idx, value in enumerate(values):
        cell = cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        add_cell_text(cell, value, center=(idx == 0))
    prevent_row_split(table.rows[-1])

# 17.4 cm usable width = 9865 DXA. Proportions closely follow the reference image.
set_table_geometry(table, [850, 2350, 3150, 3515])

doc.core_properties.title = "Bảng 3.1: Danh sách bảng"
doc.core_properties.subject = "Danh sách bảng cơ sở dữ liệu website bán kính thời trang RIOR"
doc.core_properties.author = "Nhóm dự án RIOR"
doc.core_properties.keywords = "RIOR, cơ sở dữ liệu, danh sách bảng"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
