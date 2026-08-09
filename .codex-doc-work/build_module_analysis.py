from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Java3\DA1_03")
OUT = ROOT / "Bao_cao_phan_tich_module_DA1_03.docx"
ERD = ROOT / "so_do_quan_he_thuc_the_chinh_RIOR.png"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF8E8"
GOLD = "9A6A16"
RED = "9B1C1C"
GREEN = "1E6B45"
WHITE = "FFFFFF"
BLACK = "111827"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def make_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id, fmt, text, left, hanging, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_abstract(101, "bullet", "•", 540, 270, "Symbol")
    add_abstract(102, "decimal", "%1.", 540, 270)
    add_num(201, 101)
    add_num(202, 102)
    return 201, 202


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(hp.add_run("DA1_03  |  Báo cáo phân tích module"), size=9, color=MUTED, bold=True)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("Trang "), size=9, color=MUTED)
    add_page_number(fp)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    apply_num(p, BULLET_NUM_ID)
    set_run_font(p.add_run(text))
    return p


def add_step(doc, text):
    p = doc.add_paragraph()
    apply_num(p, DECIMAL_NUM_ID)
    set_run_font(p.add_run(text))
    return p


def add_code_path(doc, label, path):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(label + ": "), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(path), name="Consolas", size=9.5, color=BLACK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        shade_cell(hdr.cells[i], LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(value)), size=9.2, color=NAVY, bold=True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2 == 1:
            for cell in cells:
                shade_cell(cell, "FAFBFC")
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, text, fill=PALE_GOLD, title_color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title + ". "), bold=True, color=title_color)
    set_run_font(p.add_run(text))
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_module_header(doc, number, title, purpose):
    doc.add_page_break()
    doc.add_heading(f"{number}. {title}", level=1)
    add_callout(doc, "Vai trò", purpose, fill=LIGHT_BLUE, title_color=NAVY)


def build():
    doc = Document()
    configure_document(doc)
    global BULLET_NUM_ID, DECIMAL_NUM_ID
    BULLET_NUM_ID, DECIMAL_NUM_ID = make_numbering(doc)

    # Cover: editorial_cover pattern.
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("BÁO CÁO KỸ THUẬT"), size=11, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("PHÂN TÍCH TOÀN BỘ MODULE"), size=28, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Dự án DA1_03 - Hệ thống quản lý bán kính"), size=16, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Logic quan trọng • Luồng xử lý • Kỹ thuật triển khai • Rủi ro và khuyến nghị"), size=10.5, color=MUTED, italic=True)
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Phạm vi: src/main, src/test và cấu hình Maven"), size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Ngày phân tích: 09/08/2026"), size=10, color=MUTED)

    doc.add_page_break()
    doc.add_heading("Tóm tắt điều hành", level=1)
    add_body(doc, "DA1_03 là ứng dụng Java Web nguyên khối (modular monolith) theo hướng MVC/layered architecture. Mã được chia theo tính năng: bán hàng tại quầy, sản phẩm, hóa đơn - thống kê, khách hàng và nhân viên. Backend dùng Servlet, Service, DAO/Repository; giao diện dùng JSP/JSTL kết hợp JavaScript; dữ liệu lưu trên SQL Server qua cả JPA/Hibernate và JDBC thuần.")
    add_callout(doc, "Kết luận chính", "Logic nghiệp vụ sâu nhất nằm ở POS và voucher: giữ tồn kho ngay khi thêm vào giỏ, khóa bi quan để tránh bán vượt tồn, tái kiểm tra voucher trước thanh toán, hoàn tồn/hoàn lượt voucher khi hủy và ghi lịch sử xuyên suốt vòng đời hóa đơn.")
    add_table(doc, ["Hạng mục", "Kết quả khảo sát", "Ý nghĩa"], [
        ["Backend", "150 tệp Java / 14.849 dòng", "Bao gồm entity, controller, service, DAO/repository và 4 lớp kiểm thử."],
        ["Frontend", "43 JSP, 10 JavaScript, 25 CSS", "Có render phía server và tương tác AJAX phong phú, đặc biệt ở màn hình POS."],
        ["Kiểm thử", "18/18 test đạt", "Xác nhận validation POS, contract DTO, route mapping và chính sách giữ lượt voucher."],
        ["Build", "Maven WAR, Java 17", "Triển khai trên Servlet container tương thích Jakarta Servlet 6, ví dụ Tomcat 10.1+."],
        ["CSDL", "SQL Server; JPA + JDBC", "JPA cho phần lớn CRUD/nghiệp vụ; JDBC dùng nhiều ở hóa đơn và thống kê."],
    ], [1800, 2700, 4860])
    add_body(doc, "Phạm vi đọc: toàn bộ mã nguồn và cấu hình trong src/, pom.xml, các tài liệu dự án và kiểm thử. Các thư mục target/ và out/ là sản phẩm build sinh tự động nên được loại khỏi phân tích nội dung để tránh đếm trùng; ảnh nhị phân chỉ được xem khi liên quan tới mô hình dữ liệu hoặc giao diện.")

    doc.add_heading("Bản đồ nội dung", level=2)
    for item in [
        "Kiến trúc, công nghệ và luồng dữ liệu tổng quát",
        "Hạ tầng dùng chung và cấu hình",
        "Module QuanLySanPham",
        "Module BanHangTaiQuay",
        "Module QuanLyHoaDon và thống kê",
        "Module QuanLyKhachHang",
        "Module QuanLyNhanVien",
        "Frontend, kiểm thử, tương tác liên module và khuyến nghị",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("1. Kiến trúc và công nghệ", level=1)
    add_table(doc, ["Lớp", "Thành phần", "Trách nhiệm"], [
        ["View", "JSP, JSTL, CSS, JavaScript", "Hiển thị dữ liệu request, thu form, gọi AJAX, lọc và phản hồi tức thời trên trình duyệt."],
        ["Controller", "Servlet / Controller", "Ánh xạ URL, đọc tham số/JSON, quản lý session, gọi service, forward JSP hoặc trả JSON."],
        ["Service", "Các interface và *ServiceImpl", "Validation, điều phối transaction, chính sách nghiệp vụ, tính tiền, tồn kho và voucher."],
        ["Persistence", "DAO/Repository", "JPQL/JPA hoặc PreparedStatement/JDBC; ánh xạ kết quả thành Entity/View Model."],
        ["Database", "SQL Server", "Lưu sản phẩm, biến thể, hóa đơn, thanh toán, lịch sử, khách hàng, nhân viên và voucher."],
    ], [1350, 2550, 5460])
    doc.add_heading("1.1. Luồng request chuẩn", level=2)
    for step in [
        "Người dùng thao tác trên JSP; form hoặc JavaScript gửi request tới URL của Servlet.",
        "Controller chuẩn hóa dữ liệu, xác định action, dựng DTO hoặc entity đầu vào.",
        "Service kiểm tra luật nghiệp vụ và mở transaction khi cần tính nhất quán.",
        "DAO/Repository truy vấn SQL Server qua JPA/Hibernate hoặc JDBC PreparedStatement.",
        "Controller trả JSP với request attribute hoặc JSON cho AJAX; JavaScript cập nhật giao diện.",
    ]:
        add_step(doc, step)

    doc.add_heading("1.2. Công nghệ chủ đạo", level=2)
    add_table(doc, ["Công nghệ", "Cách dùng trong dự án"], [
        ["Java 17 + Maven", "Biên dịch Java 17, đóng gói WAR; Surefire chạy JUnit 5."],
        ["Jakarta Servlet 6 / JSP / JSTL", "Servlet ánh xạ bằng annotation; JSP render màn quản trị và POS."],
        ["Hibernate 6.4 / JPA 3", "Entity annotation, JPQL, EntityManager, transaction và PESSIMISTIC_WRITE."],
        ["SQL Server JDBC", "PreparedStatement, CTE, OUTER APPLY, DATEPART, OFFSET/FETCH cho báo cáo và hóa đơn."],
        ["Gson", "Đọc/ghi JSON cho API POS và thống kê."],
        ["Apache POI", "Xuất Excel sản phẩm, biến thể, voucher và nhân viên."],
        ["Jakarta Mail", "Gửi thông báo đăng nhập và tài khoản nhân viên."],
        ["JavaScript thuần", "Fetch API, debounce, modal, bộ lọc, export phía client và đồng bộ dashboard."],
    ], [2500, 6860])

    if ERD.exists():
        doc.add_heading("1.3. Quan hệ thực thể giao dịch chính", level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ERD), width=Inches(6.35))
        image_shape = doc.inline_shapes[-1]
        image_shape._inline.docPr.set(
            "descr",
            "Sơ đồ quan hệ giữa khách hàng, nhân viên, hóa đơn, phiếu giảm giá, thanh toán, chi tiết hóa đơn, sản phẩm và biến thể sản phẩm.",
        )
        image_shape._inline.docPr.set("title", "Quan hệ thực thể giao dịch chính")
        doc.add_paragraph("Hình 1. Quan hệ trung tâm giữa khách hàng, nhân viên, hóa đơn, thanh toán, voucher và sản phẩm.", style="Caption")
        add_body(doc, "Hóa đơn là aggregate trung tâm: một hóa đơn có nhiều chi tiết và nhiều bản ghi thanh toán/lịch sử; mỗi chi tiết trỏ tới một biến thể sản phẩm; nhân viên lập đơn, khách hàng và voucher là liên kết tùy chọn.")

    add_module_header(doc, "2", "Hạ tầng dùng chung và cấu hình (BE, resources)", "Khởi tạo kết nối dữ liệu, EntityManager, cấu hình ứng dụng, validation, email và các tiện ích dùng chéo.")
    add_code_path(doc, "Tệp trung tâm", "pom.xml; src/main/resources/META-INF/persistence.xml; BE/jdbc; QuanLySanPham/Utils")
    doc.add_heading("2.1. Logic và kỹ thuật quan trọng", level=2)
    for text in [
        "Persistence unit 'default' quét toàn bộ entity, dùng SQLServerDialect và hbm2ddl.auto=update.",
        "EntityManagerUtlis giữ một EntityManagerFactory tĩnh và cấp EntityManager theo thao tác; mỗi service/DAO tự đóng tài nguyên.",
        "BE.DatabaseConnectionManager có thể nhận cấu hình từ system property hoặc biến môi trường; đồng thời tạo bộ JPA properties cho EntityManagerFactory.",
        "ValidationUtils chuẩn hóa khoảng trắng, kiểm tra mã, tên và liên kết bắt buộc; ValidationException mang map lỗi theo field về form.",
        "EmailService dùng SMTP STARTTLS để gửi HTML mail; được gọi khi đăng nhập hoặc khi tạo nhân viên.",
    ]:
        add_bullet(doc, text)
    add_table(doc, ["Biện pháp", "Mục đích", "Nhận xét"], [
        ["Factory + per-operation EntityManager", "Tái sử dụng metadata JPA nhưng tách persistence context cho từng thao tác", "Đúng hướng; cần đóng factory khi ứng dụng dừng."],
        ["PreparedStatement", "Truyền tham số an toàn cho JDBC", "Được dùng nhất quán trong module hóa đơn/thống kê."],
        ["Cấu hình qua môi trường", "Tách host, port, user, password khỏi code", "Chỉ BE manager hỗ trợ đầy đủ; các cấu hình khác vẫn còn giá trị cứng."],
        ["Transaction thủ công", "Commit/rollback rõ ràng", "Phù hợp app Servlet không dùng container-managed transaction."],
    ], [2350, 3820, 3190])
    add_callout(doc, "Điểm cần lưu ý", "Dự án có hai DatabaseConnectionManager và hai EntityManager utility với cách cấu hình khác nhau. persistence.xml và EmailService còn chứa bí mật trực tiếp; nên chuyển toàn bộ sang biến môi trường/secret store và thu hồi thông tin đã lộ trong lịch sử mã nguồn.", fill="FDECEC", title_color=RED)

    add_module_header(doc, "3", "Module QuanLySanPham", "Quản lý sản phẩm, biến thể, danh mục thuộc tính, ảnh, tồn kho và phiếu giảm giá; đồng thời chứa bộ entity JPA dùng chung của toàn hệ thống.")
    add_code_path(doc, "Cấu trúc", "QuanLySanPham/{controller,service,dao,Entity,dto,Utils}; webapp/Admin/QuanLySanPham; webapp/Admin/QuanLyBienThe")
    doc.add_heading("3.1. Sản phẩm và biến thể", level=2)
    add_body(doc, "SanPham là thông tin chung; SanPhamChiTiet là SKU/biến thể theo màu và kích cỡ, có mã, giá nhập, giá bán, tồn kho, trạng thái, ảnh và cờ xóa mềm. SanPham liên kết các bảng từ điển như danh mục, thương hiệu, chất liệu, kiểu dáng, tròng kính và gọng kính.")
    for text in [
        "Khi tạo sản phẩm cùng biến thể, service kiểm tra mã/tên, liên kết bắt buộc, ít nhất một biến thể và không cho trùng tổ hợp màu + kích cỡ; sản phẩm và toàn bộ biến thể được persist trong cùng transaction.",
        "Khi thêm/cập nhật biến thể riêng, service kiểm tra trùng trong request và trong database, giá/tồn/trọng lượng không âm; DAO hỗ trợ saveAll/updateAll theo lô.",
        "Tồn kho được cập nhật theo chênh lệch; hệ thống từ chối kết quả âm.",
        "Xóa sản phẩm hoặc biến thể là xóa mềm; danh sách mặc định lọc bản ghi đã xóa.",
        "Servlet dùng MultipartConfig để nhận ảnh theo màu và lưu đường dẫn phục vụ JSP/POS.",
        "Danh sách có phân trang, tìm theo tên/danh mục/thương hiệu/khoảng giá, bật/tắt trạng thái và xuất Excel bằng Apache POI.",
    ]:
        add_bullet(doc, text)
    add_table(doc, ["Thành phần", "Logic chính", "Kỹ thuật"], [
        ["SanPhamServiceImpl", "Validation tên/mã; tạo sản phẩm + biến thể nguyên tử", "Validation map; JPA transaction; Set phát hiện trùng."],
        ["SanPhamChiTietServiceImpl", "CRUD biến thể, kiểm tra tổ hợp màu-size, tồn không âm", "Batch DAO; soft delete; tìm kiếm động."],
        ["SanPhamServlet", "Form, upload ảnh, tìm kiếm, trạng thái, export", "Multipart upload; request attributes; Apache POI."],
        ["SanPhamChiTietServlet", "Quản lý SKU, tồn kho, trạng thái, export", "Route theo action; JSON cho switch trạng thái."],
    ], [2600, 4010, 2750])

    doc.add_heading("3.2. Các bảng từ điển/thuộc tính", level=2)
    add_body(doc, "DanhMuc, ThuongHieu, ChatLieu, KieuDang, MauSac, KichCo, TrongKinh, GongKinh, HinhDangGong và KieuQuaiKinh dùng cùng khuôn CRUD. Mỗi Servlet định tuyến list/new/insert/edit/update/delete; LookupServiceImpl điều phối DAO chuyên biệt.")
    add_table(doc, ["Mẫu triển khai", "Cách áp dụng"], [
        ["Generic DAO", "save/update/delete/findById/findAll/paging dùng chung qua GenericDaoImpl<T, ID>."],
        ["DAO chuyên biệt", "Bổ sung search, kiểm tra trùng hoặc quan hệ riêng cho từng loại thuộc tính."],
        ["Validation tái sử dụng", "Chuẩn hóa tên, cấm chuỗi toàn số/ký tự đặc biệt không phù hợp và kiểm tra mã."],
        ["Trạng thái thay cho xóa cứng", "Giữ khóa ngoại ổn định đối với sản phẩm đã phát sinh giao dịch."],
    ], [2600, 6760])

    doc.add_heading("3.3. Quản trị phiếu giảm giá", level=2)
    for text in [
        "Phân biệt phiếu công khai (loaiPhieu=0) và phiếu cá nhân (loaiPhieu=1), phiếu cá nhân liên kết khách hàng qua KhachHangPhieuGiamGia.",
        "Servlet tự sinh mã voucher, validate loại giảm, giá trị, số lượng, ngày bắt đầu/kết thúc, đơn tối thiểu, giảm tối đa, trạng thái và chống trùng mã.",
        "DAO hỗ trợ lọc nhiều tiêu chí, phân trang, đếm, xuất toàn bộ kết quả và cập nhật trạng thái.",
        "Excel export tạo workbook trực tiếp vào response; giao diện dùng fetch để đổi trạng thái không tải lại toàn trang.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Ranh giới module", "QuanLySanPham quản trị định nghĩa voucher; BanHangTaiQuay mới là nơi thực thi luật giữ lượt, áp dụng, tái kiểm tra và hoàn voucher trong giao dịch.")

    add_module_header(doc, "4", "Module BanHangTaiQuay (POS)", "Điều phối toàn bộ phiên bán tại quầy: hóa đơn chờ, giỏ hàng, khách hàng, voucher, thanh toán, in/hủy hóa đơn và đồng bộ tồn kho.")
    add_code_path(doc, "Tệp trung tâm", "BanHangController.java; BanHangServiceImpl.java; VoucherServiceImpl.java; BanHangDAOImpl.java; assets/js/banhang.js")
    doc.add_heading("4.1. Vòng đời hóa đơn chờ", level=2)
    for step in [
        "Lấy hoặc tạo ca làm việc đang mở cho nhân viên; mỗi nhân viên được giữ tối đa 10 hóa đơn chờ.",
        "Tạo mã hóa đơn theo thời gian, trạng thái 0, tổng tiền 0 và ghi lịch sử TAO_DON.",
        "Giao diện giữ nhiều tab hóa đơn; URL id được ưu tiên hơn session để chọn hóa đơn hiện tại.",
        "Thêm/xóa/đổi số lượng sản phẩm làm thay đổi tồn ngay và tính lại tổng tiền.",
        "Gán khách hàng hoặc chọn khách lẻ; đổi khách sẽ tự gỡ voucher cũ để tránh dùng sai quyền.",
        "Áp/revalidate voucher, xác nhận thanh toán hoặc hủy; mọi bước quan trọng đều ghi LichSuHoaDon.",
    ]:
        add_step(doc, step)

    doc.add_heading("4.2. Kiểm soát tồn kho và đồng thời", level=2)
    add_table(doc, ["Tình huống", "Logic", "Biện pháp kỹ thuật"], [
        ["Thêm vào giỏ", "Kiểm tra SKU/sản phẩm hoạt động và đủ tồn; trừ tồn ngay", "PESSIMISTIC_WRITE trên SanPhamChiTiet; transaction JPA."],
        ["Tăng/giảm số lượng", "Tính chênh lệch mới-cũ và điều chỉnh tồn ngược dấu", "Khóa SKU; từ chối nếu phần tăng vượt tồn."],
        ["Xóa dòng", "Hoàn toàn bộ số lượng của dòng về tồn", "Khóa SKU và remove ChiTietHoaDon trong cùng transaction."],
        ["Hủy hóa đơn", "Hoàn tồn mọi dòng, hoàn voucher, đặt trạng thái 5", "Khóa hóa đơn/SKU; lịch sử HUY_DON."],
        ["Thanh toán", "Không trừ tồn lần nữa; chỉ xác nhận tồn không âm", "Khóa hóa đơn và từng SKU trước khi ghi thanh toán."],
    ], [2100, 3990, 3270])
    add_callout(doc, "Ý đồ thiết kế", "Tồn kho được “giữ chỗ” ngay khi hàng vào hóa đơn chờ. Vì vậy hủy/xóa phải hoàn tồn; thanh toán không được trừ lần hai. Cách này giảm nguy cơ hai quầy cùng bán một đơn vị cuối cùng.")

    doc.add_heading("4.3. Voucher và giữ lượt", level=2)
    for text in [
        "Tìm tối đa 50 ứng viên còn hiệu lực, sau đó lọc theo tổng tiền, số lượt và quyền sở hữu; trả tối đa 8 voucher phù hợp cho UI.",
        "Khi áp: khóa hóa đơn và voucher, kiểm tra trạng thái/thời gian/số lượng/đơn tối thiểu/loại giảm; voucher cá nhân phải thuộc đúng khách hàng.",
        "Giảm phần trăm có trần giamToiDa; giảm tiền cố định bị chặn không vượt tổng hàng.",
        "soLuongDaDung được đồng bộ theo số hóa đơn đang giữ voucher, không chỉ số hóa đơn đã thanh toán.",
        "Chính sách giữ lượt theo thứ tự thời điểm AP_VOUCHER; hóa đơn áp trước được ưu tiên. Khi quá lượt, revalidate gỡ voucher của chính hóa đơn đang kiểm tra để hạn chế khóa chéo.",
        "Trước thanh toán, hệ thống revalidate: voucher hết hạn/hết lượt/không còn đủ điều kiện sẽ bị gỡ; nếu giá trị thay đổi thì tổng tiền được tính lại và trả cảnh báo cho UI.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("4.4. Thanh toán", level=2)
    for step in [
        "Chuẩn hóa mã phương thức; 'TM' được ánh xạ sang PTTT001.",
        "Tái kiểm tra voucher trong transaction riêng rồi khóa hóa đơn trong transaction thanh toán.",
        "Kiểm tra hóa đơn chưa thanh toán/chưa hủy, có sản phẩm, SKU còn kinh doanh và số lượng hợp lệ.",
        "Với tiền mặt, số tiền khách đưa phải dương và không nhỏ hơn tổng hóa đơn; với chuyển khoản có thể kiểm tra mã giao dịch không trùng.",
        "Ghi ThanhToanHoaDon và LichSuThanhToan, đặt hóa đơn trạng thái 3, lưu ngayThanhToan, cộng doanh thu ca và ghi lịch sử THANH_TOAN.",
    ]:
        add_step(doc, step)
    add_body(doc, "Scheduler HoaDonChoCuoiNgayScheduler chạy theo múi giờ Asia/Ho_Chi_Minh lúc 23:59:59, tự hủy hóa đơn trạng thái 0/1 trước mốc kết thúc ngày; khi ứng dụng dừng, executor được shutdown.")

    doc.add_heading("4.5. Controller và giao diện POS", level=2)
    add_table(doc, ["Nhóm route", "Chức năng"], [
        ["/ban-hang", "Render màn chính, hóa đơn chờ, hóa đơn hiện tại và sản phẩm."],
        ["/tao-hoa-don, /lay-hoa-don-cho", "Tạo và nạp hóa đơn chờ."],
        ["/tim-san-pham, /quet-qr", "Tìm SKU theo từ khóa hoặc mã QR."],
        ["/them-san-pham, /xoa-san-pham, /cap-nhat-so-luong", "Thao tác giỏ hàng qua JSON/AJAX."],
        ["/tra-cuu-khach-hang, /gan-khach-hang, /chon-khach-le", "Chọn/tạo/gỡ khách hàng."],
        ["/tim-voucher, /ap-voucher, /go-voucher, /revalidate-voucher", "Tìm, áp, gỡ và đồng bộ voucher."],
        ["/thanh-toan/thanh-toan, /huy-hoa-don", "Kết thúc hoặc hủy giao dịch."],
    ], [3500, 5860])
    add_body(doc, "banhang.js dùng Fetch API, debounce tìm kiếm, hàng đợi thêm sản phẩm, cập nhật giỏ từ server, quét QR bằng camera, modal khách hàng/thanh toán/hủy, tự chọn voucher tốt nhất, tính tiền thối và thông báo dashboard qua localStorage/storage event.")

    add_module_header(doc, "5", "Module QuanLyHoaDon và thống kê", "Tra cứu hóa đơn đã phát sinh, xem chi tiết/thanh toán/lịch sử, cập nhật trạng thái và tổng hợp dashboard theo thời gian.")
    add_code_path(doc, "Cấu trúc", "QuanLyHoaDon/{controller,service,dao,Model}; webapp/FE/Admin/QuanLyHoaDon; webapp/FE/Admin/Thongke.jsp")
    doc.add_heading("5.1. Quản lý hóa đơn", level=2)
    for text in [
        "Danh sách dùng JDBC join hóa đơn với nhân viên, khách hàng và voucher; phân trang SQL Server bằng OFFSET/FETCH.",
        "Bản ghi trạng thái 5 có lý do bắt đầu bằng 'Xóa mềm' được loại khỏi danh sách, nhưng hóa đơn hủy nghiệp vụ vẫn hiển thị.",
        "Chi tiết hóa đơn join sâu tới sản phẩm, biến thể và thuộc tính; OUTER APPLY chọn ảnh chính, sau đó chuẩn hóa đường dẫn ảnh.",
        "Trang chi tiết tải độc lập danh sách hàng, thanh toán, lịch sử thanh toán và lịch sử hóa đơn. Nếu một phần phụ lỗi, controller trả danh sách rỗng và vẫn hiển thị phần còn lại.",
        "Cập nhật trạng thái bằng transaction JDBC: update hoa_don và insert lich_su_hoa_don phải cùng commit hoặc cùng rollback.",
        "JavaScript phía danh sách lọc ngay trên các dòng đã tải, xuất các dòng đang thấy ra CSV và hỗ trợ in/mở modal.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("5.2. Dashboard thống kê", level=2)
    add_table(doc, ["Chỉ số", "Cách tính"], [
        ["Tổng đơn", "Đếm hóa đơn tạo trong khoảng [from, to+1 ngày)."],
        ["Doanh thu", "Tổng tong_tien_thanh_toan của hóa đơn trạng thái 3 theo ngayThanhToan, fallback ngayTao."],
        ["Sản phẩm bán", "Tổng số lượng chi tiết thuộc các hóa đơn đã thanh toán."],
        ["Hoàn tất/hủy/đang xử lý", "Phân nhóm theo trạng thái 3, 5 và các trạng thái còn lại/null."],
        ["Bán chạy/chậm", "CTE tổng bán theo sản phẩm kết hợp tổng tồn; chọn TOP 5."],
        ["Khách hàng tốt nhất", "TOP 5 theo tổng chi tiêu của hóa đơn trạng thái 3."],
        ["Chuỗi doanh thu", "Theo ngày nếu khoảng <=62 ngày; theo tháng nếu dài hơn; service lấp điểm thiếu bằng 0."],
    ], [2700, 6660])
    add_body(doc, "ThongKeController vừa render JSP vừa cung cấp JSON action=overview và action=revenue-series. Giao diện chọn ngày/tháng/quý/năm, vẽ biểu đồ bằng DOM/SVG/CSS, xuất báo cáo và tự tải lại khi POS phát tín hiệu thay đổi số liệu.")

    add_module_header(doc, "6", "Module QuanLyKhachHang", "Quản lý hồ sơ khách hàng, tra cứu cho POS, địa chỉ nhận hàng và địa chỉ mặc định.")
    add_code_path(doc, "Cấu trúc", "QuanLyKhachHang/{servlet,repository}; webapp/QuanLyKhachHang; Entity KhachHang và DiaChiKhachHang")
    for text in [
        "Danh sách chỉ lấy khách đang hoạt động; tìm kiếm POS ưu tiên trùng chính xác/prefix số điện thoại hoặc mã trước kết quả chứa từ khóa, giới hạn 8 bản ghi.",
        "Nếu mã khách trống, repository sinh KHxxx bằng truy vấn MAX kết hợp SQL Server UPDLOCK + HOLDLOCK để giảm trùng mã khi tạo đồng thời.",
        "Cập nhật khách chỉ thay các trường cho phép và đồng bộ tên/SĐT người nhận sang toàn bộ địa chỉ trong cùng transaction.",
        "Đổi trạng thái là active/inactive, không xóa vật lý.",
        "Thêm địa chỉ lấy tên và SĐT từ khách hàng; tỉnh/phường dùng mã từ API hành chính trên giao diện; trường quận/huyện được giữ rỗng để đáp ứng schema hiện tại.",
        "Đặt mặc định duyệt toàn bộ địa chỉ của khách, đặt về 0 rồi bật địa chỉ chọn lên 1 trong một transaction.",
    ]:
        add_bullet(doc, text)
    add_table(doc, ["Luồng", "Route", "Kỹ thuật"], [
        ["Danh sách/thêm/sửa/khóa", "/khach-hang/*", "Servlet switch theo path; flash message trong session; JPA repository."],
        ["Địa chỉ", "/dia-chi-khach-hang/*", "Quan hệ ManyToOne; transaction đặt một địa chỉ mặc định."],
        ["Tra cứu tại POS", "Gọi qua BanHangService", "JPQL ranking bằng CASE; tạo mới nếu số điện thoại chưa tồn tại."],
        ["Chọn tỉnh/phường", "provinces.open-api.vn", "Fetch API phía trình duyệt; lưu code và tên hiển thị."],
    ], [2500, 3000, 3860])

    add_module_header(doc, "7", "Module QuanLyNhanVien và xác thực", "Quản lý tài khoản nhân viên, vai trò/chức vụ, phân trang, tìm kiếm, xuất Excel, đăng ký, đăng nhập và session.")
    add_code_path(doc, "Cấu trúc", "QuanLyNhanVien/{controller,service}; QuanLySanPham/dao/NhanVienDao; LoginServlet; RegisterServlet")
    for text in [
        "Danh sách chỉ lấy nhân viên trạng thái 1, phân trang 10 bản ghi; tìm theo họ tên, mã hoặc email.",
        "Mã nhân viên sinh dạng NV0001 từ mã lớn nhất; mật khẩu tạo ngẫu nhiên 8 ký tự bằng SecureRandom khi mở form thêm.",
        "Chức vụ bị giới hạn trong ba giá trị hợp lệ; xóa là đặt trạng thái 0.",
        "Tạo nhân viên gửi email nền chứa mã, email, mật khẩu và chức vụ; export Excel dùng Apache POI với style, merge title và độ rộng cột.",
        "Đăng nhập nhận email hoặc mã nhân viên, kiểm tra tài khoản hoạt động, lưu NhanVien vào session và chuyển quản lý tới dashboard, nhân viên tới sản phẩm.",
        "Đăng ký công khai luôn gán vai trò nhân viên thường, tránh tự cấp quyền quản lý.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Lỗi logic cần ưu tiên", "RegisterServlet lưu PasswordUtil.hash(mật khẩu) nhưng NhanVienServiceImpl.dangNhap lại so sánh chuỗi người dùng nhập trực tiếp với nv.matKhau. Trong khi luồng quản trị nhân viên lưu/gửi mật khẩu thô. Hai chiến lược không đồng nhất có thể khiến tài khoản tự đăng ký không đăng nhập được và làm giảm an toàn tài khoản.", fill="FDECEC", title_color=RED)

    add_module_header(doc, "8", "Frontend và trải nghiệm người dùng", "Kết nối dữ liệu backend với các màn quản trị/POS, cung cấp lọc, modal, phản hồi AJAX, in và xuất dữ liệu.")
    add_table(doc, ["Khu vực", "Logic phía trình duyệt"], [
        ["POS", "Fetch JSON; debounce; queue thêm SKU; QR camera; modal thanh toán; tiền thối; voucher tốt nhất; đồng bộ dashboard."],
        ["Sản phẩm/biến thể", "Switch trạng thái bằng fetch; form nhiều biến thể; preview/upload ảnh; lọc và export."],
        ["Hóa đơn", "Lọc theo từ khóa/loại/trạng thái/ngày trên dòng hiện có; CSV; in; tab chi tiết."],
        ["Thống kê", "Chọn khoảng nhanh hoặc ngày tùy chỉnh; gọi JSON; cập nhật card/đồ thị; export."],
        ["Khách hàng", "Validate client, modal form, lọc tức thời; địa chỉ gọi API tỉnh/phường."],
        ["Nhân viên/voucher", "Validate form, AJAX đổi trạng thái/xóa, toast, autocomplete và export."],
    ], [2550, 6810])
    add_body(doc, "Các layout sidebar/header được include lại giữa các JSP. Tuy nhiên hiện tồn tại nhiều cây giao diện gần trùng nhau (Admin, FE/Admin và Admin/Admin), cùng các bản CSS/JS lặp; đây là dấu hiệu lịch sử chuyển cấu trúc và làm tăng nguy cơ sửa một bản nhưng chạy bản khác.")

    add_module_header(doc, "9", "Kiểm thử và mức độ xác nhận", "Đánh giá những gì đã có test tự động và phần logic nào vẫn chỉ được xác nhận bằng đọc mã.")
    add_table(doc, ["Test", "Số ca", "Phạm vi"], [
        ["BanHangServiceValidationTest", "9", "ID/số lượng không hợp lệ, ca làm việc và giới hạn 10 hóa đơn chờ."],
        ["DtoContractTest", "3", "Getter/setter các request DTO và kết quả revalidate voucher."],
        ["RouteMappingTest", "2", "Các route chính của POS và thanh toán."],
        ["VoucherReservationPolicyTest", "4", "Giữ lượt, trả lượt, thứ tự ưu tiên và slot chờ."],
    ], [3000, 950, 5410])
    add_callout(doc, "Kết quả chạy", "Ngày 09/08/2026, Maven test chạy 18 ca: 0 failure, 0 error, 0 skipped; BUILD SUCCESS. Maven đồng thời cảnh báo dependency jakarta.mail bị khai báo trùng.", fill="EAF6EF", title_color=GREEN)
    add_body(doc, "Khoảng trống kiểm thử: chưa có integration test đầy đủ với database cho thanh toán/hoàn tồn/voucher, chưa có test controller/JSP, chưa có test bảo mật/phân quyền, chưa kiểm tra upload/export và chưa có kịch bản đồng thời nhiều quầy. Do hbm2ddl.auto=update, việc khởi tạo JPA trong test còn có thể thay đổi schema database cấu hình.")

    add_module_header(doc, "10", "Tương tác liên module", "Cho thấy các module không độc lập hoàn toàn và xác định điểm nối cần chú ý khi sửa tính năng.")
    add_table(doc, ["Nguồn", "Phụ thuộc", "Mục đích"], [
        ["BanHangTaiQuay", "QuanLySanPham Entity/Service", "SKU, giá, tồn, hóa đơn, chi tiết, voucher, payment."],
        ["BanHangTaiQuay", "QuanLyNhanVien", "Nhân viên/session và ca làm việc."],
        ["QuanLyHoaDon", "SQL Server schema chung", "Đọc dữ liệu giao dịch do POS ghi bằng JDBC."],
        ["QuanLyKhachHang", "QuanLySanPham.Entity", "Dùng entity khách hàng và địa chỉ chung."],
        ["QuanLyNhanVien", "QuanLySanPham.dao/Entity/Utils", "DAO nhân viên, entity và email."],
        ["Dashboard", "POS frontend", "POS phát localStorage event để dashboard làm mới số liệu."],
    ], [2400, 2900, 4060])
    doc.add_heading("10.1. Ví dụ truy vết thay đổi", level=2)
    add_table(doc, ["Muốn sửa", "Nên đi theo chuỗi tệp"], [
        ["Luật thanh toán POS", "banhang.js → BanHangController → BanHangServiceImpl → Entity thanh toán/lịch sử → HoaDonDAO màn quản lý."],
        ["Voucher", "phieu-giam-gia-form.jsp/Servlet/DAO → VoucherServiceImpl → BanHangServiceImpl → lịch sử hóa đơn."],
        ["Tồn kho", "SanPhamChiTietService/DAO → POS add/update/delete/cancel/payment → thống kê sản phẩm."],
        ["Dashboard", "Thongke.jsp/thongke.js → ThongKeController/Service/DAO → trạng thái/ngày thanh toán POS."],
        ["Khách hàng", "customer JSP/Servlet/Repository → POS tra cứu/gán khách → voucher cá nhân → hóa đơn."],
    ], [2700, 6660])

    add_module_header(doc, "11", "Rủi ro và khuyến nghị", "Ưu tiên các thay đổi giúp hệ thống an toàn, nhất quán và dễ bảo trì hơn mà không làm mất logic nghiệp vụ hiện có.")
    add_table(doc, ["Ưu tiên", "Vấn đề", "Khuyến nghị"], [
        ["P0", "Thông tin DB/SMTP nằm trực tiếp trong source", "Thu hồi/đổi secret; dùng biến môi trường hoặc secret manager; xóa khỏi lịch sử Git nếu cần."],
        ["P0", "Mật khẩu hash và so sánh không đồng nhất", "Dùng một PasswordEncoder có salt (Argon2/bcrypt/PBKDF2); migrate dữ liệu; không gửi/lưu mật khẩu thô."],
        ["P0", "Không thấy Filter bảo vệ route trong src", "Thêm AuthenticationFilter và AuthorizationFilter; chặn theo vai trò ở server, không chỉ ẩn menu."],
        ["P1", "Hai hệ kết nối DB/JPA khác nhau", "Hợp nhất một cấu hình env-aware; thống nhất transaction và lifecycle EntityManagerFactory."],
        ["P1", "hbm2ddl.auto=update và load script thiếu kiểm soát", "Dùng Flyway/Liquibase; production để validate/none; database test riêng."],
        ["P1", "Ngoại lệ bị nuốt ở một số repository", "Log có ngữ cảnh và trả lỗi lên service/controller; không giả thành danh sách rỗng nếu là lỗi hệ thống."],
        ["P1", "Đặt địa chỉ mặc định chưa kiểm tra địa chỉ thuộc khách", "Query theo cả idDiaChi và idKhachHang hoặc kiểm tra quan hệ trước khi set."],
        ["P2", "Magic number trạng thái 0/1/3/5", "Dùng enum/value object và mapper; tập trung transition hợp lệ."],
        ["P2", "Cây JSP/CSS/JS trùng", "Chọn một cây chuẩn, xóa bản legacy sau kiểm chứng, đóng gói component/layout dùng chung."],
        ["P2", "Scheduler chạy trong từng instance web", "Nếu scale nhiều node, dùng job lock trong DB hoặc scheduler tập trung."],
        ["P2", "Test thiên về validation đơn vị", "Bổ sung Testcontainers/integration test cho checkout, rollback, concurrent stock và voucher."],
        ["P3", "pom.xml khai báo mail trùng", "Giữ một dependency và thêm build lint/Enforcer."],
    ], [900, 3520, 4940], font_size=8.7)

    doc.add_heading("11.1. Lộ trình đề xuất", level=2)
    for step in [
        "Tuần 1: xoay vòng secret, thống nhất password hashing, thêm filter xác thực/phân quyền và test đăng nhập.",
        "Tuần 2: hợp nhất cấu hình database; chuyển schema sang migration; tạo profile database test riêng.",
        "Tuần 3: viết integration test cho POS: giữ tồn, hoàn tồn, thanh toán, voucher hết lượt và rollback.",
        "Tuần 4: chuẩn hóa enum trạng thái, error handling và log; dọn cây frontend trùng sau khi lập route map.",
    ]:
        add_step(doc, step)

    doc.add_page_break()
    doc.add_heading("Phụ lục A. Chỉ mục module và tệp tiêu biểu", level=1)
    add_table(doc, ["Module", "Tệp nên đọc trước", "Điểm vào giao diện"], [
        ["BE / cấu hình", "persistence.xml; DatabaseConnectionManager; EntityManagerUtlis", "Login.jsp; WEB-INF/web.xml"],
        ["QuanLySanPham", "SanPhamServiceImpl; SanPhamChiTietServiceImpl; LookupServiceImpl", "Admin/QuanLySanPham; Admin/QuanLyBienThe"],
        ["BanHangTaiQuay", "BanHangController; BanHangServiceImpl; VoucherServiceImpl", "Admin/BanHangTaiQuay/ban-hang.jsp; assets/js/banhang.js"],
        ["QuanLyHoaDon", "HoaDonController/DAO; ThongKeController/Service/DAO", "FE/Admin/QuanLyHoaDon; FE/Admin/Thongke.jsp"],
        ["QuanLyKhachHang", "KhachHangServlet/Repository; DiaChiKhachHangServlet/Repository", "QuanLyKhachHang/*.jsp"],
        ["QuanLyNhanVien", "NhanVienServlet; NhanVienServiceImpl; NhanVienDaoImpl", "Admin/QuanLyNhanVien/*.jsp"],
        ["Test", "4 lớp trong src/test/java/BanHangTaiQuay/Test", "Chạy: mvnw.cmd test"],
    ], [1950, 4300, 3110], font_size=8.8)

    doc.add_heading("Phụ lục B. Quy ước trạng thái suy ra từ mã", level=1)
    add_table(doc, ["Giá trị", "Ý nghĩa đang dùng", "Nơi thể hiện"], [
        ["0", "Hóa đơn mới/chờ xác nhận; record inactive ở một số entity khác", "Tạo hóa đơn POS; lọc trạng thái theo từng entity."],
        ["1", "Hóa đơn còn chờ hoặc bản ghi đang hoạt động", "POS cho sửa trạng thái 0/1; sản phẩm/voucher/nhân viên active."],
        ["3", "Hóa đơn đã thanh toán", "Checkout, doanh thu và thống kê chỉ tính trạng thái 3."],
        ["5", "Hóa đơn đã hủy hoặc xóa mềm theo lý do", "Hủy POS; quản lý hóa đơn lọc theo ly_do_huy."],
    ], [1200, 3900, 4260])
    add_callout(doc, "Lưu ý", "Các giá trị trạng thái được diễn giải theo cách mã hiện tại kiểm tra. Vì cùng một số có ý nghĩa khác nhau giữa entity, nên enum riêng cho từng domain là cần thiết.")

    doc.core_properties.title = "Phân tích toàn bộ module dự án DA1_03"
    doc.core_properties.subject = "Logic quan trọng và kỹ thuật triển khai"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.keywords = "DA1_03, Java, Servlet, JSP, POS, module, phân tích"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
